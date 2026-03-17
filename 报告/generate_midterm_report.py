#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成中期报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = OxmlElement(f'w:{edge}')
            for key in ('sz', 'val', 'color', 'space'):
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_report():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ========== 封面 ==========
    # 标题
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('本科毕业论文（设计）中期报告')
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # 题目
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('题 目：公交准点率分析系统设计与实现')
    run.font.size = Pt(16)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = True

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # 信息表
    info_items = [
        ('专    业', '软件工程'),
        ('学    生', '吕同杰'),
        ('学    号', '2022211844'),
        ('指导教师', '韩希先'),
        ('日    期', '2026.3.17'),
    ]

    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}       {value}       ')
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_paragraph()
    doc.add_paragraph()

    # 学校
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('哈尔滨工业大学教务处制')
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = True

    # 分页
    doc.add_page_break()

    # ========== 正文部分 ==========
    # 辅助函数：添加标题段落
    def add_section_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        return p

    def add_body_text(text, indent=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        return p

    def add_sub_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = True
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.first_line_indent = Cm(0.74)
        return p

    # ========== 一、论文工作是否按开题报告预定的内容及进度安排进行 ==========
    add_section_title('一、论文（设计）工作是否按开题报告预定的内容及进度安排进行')

    add_body_text(
        '本毕业设计工作基本按照开题报告预定的内容及进度安排进行，整体进展顺利。'
        '根据开题报告中的进度规划，项目分为四个阶段：第一阶段为文献调研与需求分析（第1-3周），'
        '第二阶段为系统设计与开发（第4-11周），第三阶段为系统测试与优化（第12-14周），'
        '第四阶段为论文撰写与答辩（第15-16周）。'
    )

    add_body_text(
        '目前项目处于第二阶段的后期，系统的核心功能模块已基本开发完成。'
        '在技术选型方面，经过实际开发中的评估和调整，后端框架由开题报告中规划的Django调整为Flask，'
        '主要原因是Flask更加轻量灵活，适合本系统以数据处理和API服务为核心的特点，'
        '能够更高效地实现RESTful API接口。数据库方面，采用PostgreSQL作为主要存储方案，'
        '与开题报告中的规划一致。前端技术栈完全按照规划采用Vue 3 + Vite + Element Plus的组合，'
        '地图可视化使用Leaflet实现。'
    )

    add_body_text(
        '在功能实现方面，开题报告中规划的数据获取与导入、数据预处理、准点率计算引擎、'
        'RESTful API服务、前端可视化界面等核心模块均已完成开发。此外，还额外完成了多地区数据支持'
        '（旧金山、纽约、悉尼）和用户认证系统等扩展功能，超出了原定计划的范围。'
        '总体而言，项目进度略超前于预期。'
    )

    # ========== 二、已完成的研究工作及成果 ==========
    add_section_title('二、已完成的研究工作及成果')

    add_body_text('截至目前，已完成的研究工作及成果主要包括以下几个方面：')

    # 1. 系统总体架构
    add_sub_title('1. 系统总体架构设计与实现')
    add_body_text(
        '完成了系统的分层架构设计，自下而上包括数据层、处理层、服务层和表现层四个逻辑层次。'
        '后端采用Python 3 + Flask框架，前端采用Vue 3 + Vite + Element Plus，'
        '数据库采用PostgreSQL，整体架构职责清晰、耦合度低。系统采用前后端分离的开发模式，'
        '后端通过RESTful API提供数据服务，前端通过Axios进行HTTP请求，'
        '使用Pinia进行全局状态管理，Vue Router实现页面路由。'
    )

    # 2. 多源GTFS数据获取与导入
    add_sub_title('2. 多源GTFS数据获取与导入模块')
    add_body_text(
        '开发了三个独立的数据获取器，分别对接旧金山511 SF Bay API、纽约MTA Open Data和'
        '悉尼TfNSW Open Data三个数据源。实现了GTFS静态数据（ZIP格式）的自动解析和'
        '批量导入PostgreSQL功能，支持按地区参数导入，自动处理表依赖关系，'
        '批量插入优化（每批1000条），并具备数据验证能力。同时实现了GTFS Realtime实时数据流的解析，'
        '支持车辆位置（Vehicle Position）、行程更新（Trip Update）和服务警报（Service Alert）'
        '三种Feed类型的Protobuf数据解码。'
    )

    # 3. 数据库设计
    add_sub_title('3. 数据库设计与优化')
    add_body_text(
        '设计并实现了完整的PostgreSQL数据库方案，包含GTFS标准表（agency、routes、stops、'
        'trips、stop_times、calendar、calendar_dates、shapes等）、SF Muni扩展表'
        '（route_attributes、directions等）、准点率分析表（realtime_delay_records、'
        'route_daily_punctuality、stop_daily_punctuality、hourly_punctuality_stats等）'
        '以及用户认证表。所有表均支持region字段实现多地区数据隔离，'
        '并为常用查询字段创建了索引优化，包括线路查询索引、站点位置索引、班次查询索引和时刻表查询索引。'
    )

    # 4. 速度计算
    add_sub_title('4. 车辆速度计算模块')
    add_body_text(
        '实现了基于Haversine公式的车辆实时速度计算功能，能够根据连续GPS位置数据计算车辆瞬时速度。'
        '模块内置了GPS异常过滤机制，包括最大速度限制（120 km/h）、最小距离阈值（5米）和'
        '最小时间间隔控制（5秒），有效处理了GPS信号漂移和定位异常问题。'
        '该模块维护车辆位置历史记录，通过计算相邻两个GPS点之间的地球表面距离和时间差，'
        '得出车辆的实时运行速度。'
    )

    # 5. 准点率计算引擎
    add_sub_title('5. 准点率计算引擎')
    add_body_text(
        '开发了完整的准点率计算与分析模块（PunctualityCalculator），实现了延误记录管理、'
        '准点状态判断（提前、准点、延误、严重延误四级分类）、线路准点率统计、站点准点率统计、'
        '时段准点率统计（区分早高峰7-9点和晚高峰17-19点）以及系统准点率概览等功能。'
        '准点判断采用可配置的阈值体系：早于计划时间60秒判定为提前，延误不超过120秒判定为准点，'
        '延误120至300秒判定为延误，超过300秒判定为严重延误。'
        '支持按线路、时段、站点等多个维度的灵活计算，并提供数据导出能力。'
    )

    # 6. RESTful API
    add_sub_title('6. RESTful API服务')
    add_body_text(
        '基于Flask框架开发了30余个RESTful API接口，涵盖健康检查、地区管理、运营机构查询、'
        '线路查询（支持分页、搜索、筛选）、站点查询（支持地理位置筛选）、班次查询、'
        '线路轨迹查询、服务日历查询、实时车辆位置、实时延误信息、'
        '准点率分析（概览、线路、站点、时段）以及用户认证（登录、注册、登出）等功能。'
        '所有接口支持region参数过滤，采用统一的JSON响应格式，具备完善的错误处理和CORS跨域支持。'
        '主要接口包括：/api/routes（线路）、/api/stops（站点）、/api/trips（班次）、'
        '/api/realtime/vehicles（实时车辆）、/api/punctuality/overview（准点率概览）等。'
    )

    # 7. 前端可视化
    add_sub_title('7. 前端可视化应用')
    add_body_text(
        '开发了基于Vue 3 Composition API的单页面应用，包含11个页面视图：'
        '登录页（Login）、首页数据统计仪表盘（Home）、线路列表页（Routes）、'
        '线路详情页（RouteDetail）、站点列表页（Stops）、站点详情页（StopDetail）、'
        '地图视图页（Map）、准点率概览页（PunctualityOverview）、'
        '线路准点率分析页（RoutePunctuality）、站点准点率分析页（StopPunctuality）'
        '和实时监控页（RealtimeMonitor）。使用Pinia进行状态管理（6个Store模块），'
        'Axios封装API请求层，Leaflet实现地图可视化展示线路轨迹和站点位置。'
        '前端具备搜索筛选、分页浏览、响应式布局、自动数据刷新等特性。'
    )

    # 8. 多地区支持
    add_sub_title('8. 多地区数据支持')
    add_body_text(
        '实现了旧金山、纽约、悉尼三个城市的公交数据支持。通过数据库region字段和'
        '前端地区选择器（RegionSelector组件）实现数据隔离和切换，'
        '所有API接口和前端页面均支持多地区数据的查询和展示。'
        '后端为每个地区开发了独立的数据获取器：GTFSDataFetcher（旧金山511 API）、'
        'MTADataFetcher（纽约MTA API）、TfNSWDataFetcher（悉尼TfNSW API），'
        '能够分别从各自的开放数据平台获取GTFS静态数据和实时数据。'
        '数据导入工具支持--region参数，可按地区独立导入和管理数据。'
    )

    # 9. 用户认证
    add_sub_title('9. 用户认证系统')
    add_body_text(
        '实现了基于Token的用户认证机制，包括用户注册、登录、登出功能。'
        '密码采用PBKDF2-HMAC-SHA256算法进行哈希存储，Token使用secrets模块生成安全随机令牌。'
        '前端通过路由守卫实现页面访问控制，未登录用户自动跳转至登录页面。'
        '系统首次启动时自动创建默认管理员账户，方便开发和测试使用。'
    )

    # ========== 三、后期拟完成的研究工作及进度安排 ==========
    add_section_title('三、后期拟完成的研究工作及进度安排')

    add_body_text('后期拟完成的研究工作主要集中在系统测试优化和论文撰写两个方面，具体安排如下：')

    add_sub_title('1. 系统集成测试与数据验证（第12周）')
    add_body_text(
        '对系统所有模块进行全面的集成测试，使用旧金山、纽约、悉尼三个城市的真实GTFS数据'
        '进行端到端测试。重点验证数据导入的完整性和准确性、准点率计算结果的正确性、'
        'API接口的稳定性和响应速度，以及前端页面的功能完整性和交互体验。'
    )

    add_sub_title('2. 性能优化与问题修复（第13周）')
    add_body_text(
        '针对测试中发现的问题进行修复和优化。重点包括：数据库查询性能优化'
        '（复杂联表查询的执行效率）、大数据量场景下的分页和加载优化、'
        '前端地图渲染性能优化（大量站点和轨迹数据的展示）、'
        '实时数据刷新机制的稳定性优化。'
    )

    add_sub_title('3. 数据可视化增强（第13-14周）')
    add_body_text(
        '进一步完善前端可视化功能，增加准点率趋势图表（基于ECharts）、'
        '延误分布直方图、高峰时段对比分析图等统计图表，'
        '提升数据分析的直观性和可读性。'
    )

    add_sub_title('4. 论文撰写（第14-16周）')
    add_body_text(
        '整理研究成果，撰写毕业论文。论文将涵盖系统需求分析、总体架构设计、'
        '关键技术实现（GTFS数据解析、准点率计算算法、多源数据融合）、'
        '系统测试与结果分析等内容。同时准备答辩PPT和演示材料。'
    )

    # ========== 四、存在的问题与困难 ==========
    add_section_title('四、存在的问题与困难')

    add_body_text('在项目开发过程中，遇到了以下主要问题与困难：')

    add_sub_title('1. GTFS Realtime数据获取的网络限制')
    add_body_text(
        '部分数据源（如511 SF Bay API）的实时数据接口在国内网络环境下访问不稳定，'
        '需要通过VPN等方式解决网络连通性问题。这对实时数据的持续采集和准点率分析的时效性'
        '造成了一定影响。目前通过配置代理和增加请求重试机制来缓解该问题。'
    )

    add_sub_title('2. 多源异构数据的标准化处理')
    add_body_text(
        '虽然三个城市的数据均遵循GTFS规范，但在实际数据中存在字段差异。'
        '例如纽约MTA数据包含route_sort_order等非标准字段，'
        '悉尼TfNSW数据的部分字段命名与标准略有不同。'
        '为此在数据导入模块中实现了自动跳过数据库中不存在列的容错机制，'
        '但仍需要针对不同数据源进行适配处理。'
    )

    add_sub_title('3. 准点率计算的精度问题')
    add_body_text(
        '准点率计算依赖于实时车辆位置数据与计划时刻表的匹配，'
        '但GPS定位存在漂移误差，且不同数据源的时间戳精度不一致。'
        '目前已实现基于速度阈值和距离阈值的GPS异常过滤，'
        '但在到站时间推断的精度方面仍有提升空间。'
    )

    add_sub_title('4. 前端地图性能优化')
    add_body_text(
        '当同时展示大量站点标记和线路轨迹时，Leaflet地图的渲染性能会出现明显下降。'
        '后续需要引入标记聚合、视口裁剪等优化策略来改善大数据量场景下的地图交互体验。'
    )

    # ========== 五、论文按时完成的可能性 ==========
    add_section_title('五、论文按时完成的可能性')

    add_body_text(
        '综合当前的项目进展情况，论文按时完成的可能性很大。'
    )

    add_body_text(
        '从系统开发角度来看，核心功能模块（数据获取与导入、准点率计算引擎、'
        'RESTful API服务、前端可视化应用）已全部完成开发，'
        '系统已具备完整的端到端功能链路。多地区数据支持和用户认证系统等扩展功能也已实现，'
        '整体开发进度略超前于开题报告中的计划安排。'
    )

    add_body_text(
        '从后续工作量来看，剩余工作主要集中在系统测试优化、数据可视化增强和论文撰写三个方面。'
        '系统测试和优化工作的范围明确，预计可在2周内完成。'
        '论文撰写方面，系统架构设计文档、API接口文档等技术文档在开发过程中已同步编写，'
        '为论文撰写积累了充分的素材。'
    )

    add_body_text(
        '因此，按照当前的进度和工作计划，有充分的信心在规定时间内完成毕业论文的撰写和系统的最终交付。'
    )

    # 保存文件
    output_path = os.path.join(os.path.dirname(__file__), '吕同杰-中期报告.docx')
    doc.save(output_path)
    print(f'已生成: {output_path}')
    return output_path

if __name__ == '__main__':
    create_report()
